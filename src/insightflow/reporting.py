from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from jinja2 import Template

from .ai_assistant import executive_summary
from .diagnostics import DiagnosticIssue


HTML_TEMPLATE = Template(
    """
<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">
<title>InsightFlow 经营分析报告</title>
<style>
body{font-family:Arial,"Microsoft YaHei",sans-serif;background:#f5f7fb;color:#182235;margin:0}
.container{max-width:1180px;margin:0 auto;padding:30px}.hero{background:#17243d;color:white;border-radius:18px;padding:28px}
.kpis{display:grid;grid-template-columns:repeat(auto-fit,minmax(170px,1fr));gap:14px;margin:20px 0}
.card{background:white;border-radius:14px;padding:18px;box-shadow:0 4px 18px rgba(20,34,70,.06)}
.label{font-size:13px;color:#68738a}.value{font-size:27px;font-weight:700;margin-top:8px}
section{margin-top:24px}.issue{border-left:5px solid #7b879d}.P0,.P1{border-left-color:#c23632}.P2{border-left-color:#d17d00}.P3{border-left-color:#287d4b}
table{width:100%;border-collapse:collapse;background:white;font-size:13px}th,td{padding:9px;border-bottom:1px solid #e9edf4;text-align:left}th{background:#f0f3f8}
.small{font-size:12px;color:#aab3c2}.summary{white-space:pre-wrap;line-height:1.7}.grid{display:grid;grid-template-columns:1fr 1fr;gap:18px}
@media(max-width:800px){.grid{grid-template-columns:1fr}}
</style></head><body><div class="container">
<div class="hero"><h1>InsightFlow 经营分析与决策报告</h1>
<p class="small">分析期间：{{ period }}｜生成时间：{{ generated_at }}｜摘要引擎：{{ summary_engine }}</p>
<p class="small">数据状态：{{ data_profile }}</p>
<div class="summary">{{ summary }}</div></div>
<div class="kpis">{% for label,value in kpis %}<div class="card"><div class="label">{{ label }}</div><div class="value">{{ value }}</div></div>{% endfor %}</div>
<section><h2>关键诊断与行动</h2>{% for issue in issues %}<div class="card issue {{ issue.severity }}"><b>{{ issue.severity }} · {{ issue.title }}</b><p>{{ issue.finding }}</p><p><b>证据：</b>{{ issue.evidence }}</p><p><b>建议：</b>{{ issue.recommendation }}</p></div>{% endfor %}</section>
<div class="grid"><section><h2>收入变化驱动</h2>{{ drivers_table }}</section><section><h2>利润与成本结构</h2>{{ profit_table }}</section></div>
<div class="grid"><section><h2>目标状态</h2>{{ targets_table }}</section><section><h2>未来预测</h2>{{ forecast_table }}</section></div>
<div class="grid"><section><h2>库存风险</h2>{{ inventory_table }}</section><section><h2>数据质量</h2>{{ quality_table }}</section></div>
<section><h2>月度趋势</h2>{{ trend_table }}</section>
</div></body></html>
"""
)


def format_kpis(kpi: dict[str, float]) -> list[tuple[str, str]]:
    return [
        ("净销售额", f"£{kpi['revenue']:,.0f}"),
        ("毛利润", f"£{kpi.get('gross_profit', 0):,.0f}"),
        ("毛利率", f"{kpi.get('gross_margin', 0):.1%}"),
        ("贡献利润", f"£{kpi.get('contribution_profit', 0):,.0f}"),
        ("贡献利润率", f"{kpi.get('contribution_margin', 0):.1%}"),
        ("订单量", f"{kpi['orders']:,.0f}"),
        ("活跃客户", f"{kpi['active_customers']:,.0f}"),
        ("复购率", f"{kpi['repeat_rate']:.1%}"),
        ("取消率", f"{kpi['cancellation_rate']:.1%}"),
    ]


def present_drivers(drivers: pd.DataFrame) -> pd.DataFrame:
    labels = {
        "driver": "驱动因素",
        "previous": "上期",
        "current": "本期",
        "change": "变化",
        "revenue_contribution": "净销售额贡献",
    }
    columns = [c for c in labels if c in drivers.columns]
    return drivers[columns].rename(columns=labels).copy()



def present_targets(df: pd.DataFrame | None) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    labels = {
        "target_period": "期间",
        "metric_id": "指标",
        "actual_value": "实际值",
        "target_value": "目标值",
        "attainment": "完成度",
        "status": "状态",
        "owner": "负责人",
    }
    metrics = {
        "net_revenue": "净销售额",
        "contribution_profit": "贡献利润",
        "gross_margin": "毛利率",
        "cancellation_rate": "取消率",
    }
    out = df[[c for c in labels if c in df.columns]].rename(columns=labels).copy()
    if "指标" in out:
        out["指标"] = out["指标"].map(metrics).fillna(out["指标"])
    if "完成度" in out:
        out["完成度"] = out["完成度"].map(lambda x: f"{x:.1%}" if pd.notna(x) else "")
    return out


def present_forecast(df: pd.DataFrame | None) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    labels = {"month": "月份", "forecast": "预测值", "lower": "区间下限", "upper": "区间上限"}
    return df[[c for c in labels if c in df.columns]].rename(columns=labels).copy()


def present_inventory(df: pd.DataFrame | None) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    labels = {
        "stock_code": "SKU",
        "description": "商品",
        "category": "品类",
        "supplier": "供应商",
        "supplier_lead_days": "提前期",
        "inventory_on_hand": "现有库存",
        "avg_daily_units": "日均销量",
        "days_of_supply": "可售天数",
        "inventory_value": "库存价值",
    }
    return df[[c for c in labels if c in df.columns]].rename(columns=labels).copy()


def present_quality(df: pd.DataFrame | None) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    labels = {"dimension": "质量维度", "score_100": "得分", "description": "说明"}
    return df[[c for c in labels if c in df.columns]].rename(columns=labels).copy()

def _safe_table(df: pd.DataFrame | None, empty_text: str = "暂无数据") -> str:
    if df is None or df.empty:
        return f"<div class='card'>{empty_text}</div>"
    return df.round(3).to_html(index=False, border=0)


def build_markdown_report(
    period: str,
    kpi: dict[str, float],
    issues: list[DiagnosticIssue],
    drivers: pd.DataFrame,
    *,
    targets: pd.DataFrame | None = None,
    forecast: pd.DataFrame | None = None,
    data_profile: str | None = None,
) -> str:
    summary, engine = executive_summary(issues)
    lines = [
        "# InsightFlow 经营周报",
        "",
        f"- 分析期间：{period}",
        f"- 生成时间：{datetime.now():%Y-%m-%d %H:%M}",
        f"- 摘要引擎：{engine}",
        f"- 数据状态：{data_profile or '未提供来源标签'}",
        "",
        "## 管理层摘要",
        "",
        summary,
        "",
        "## 核心指标",
        "",
    ]
    for label, value in format_kpis(kpi):
        lines.append(f"- **{label}**：{value}")
    lines.extend(["", "## 关键诊断", ""])
    for issue in issues:
        lines.extend(
            [
                f"### {issue.severity}｜{issue.title}",
                "",
                issue.finding,
                "",
                f"**证据：** {issue.evidence}",
                "",
                f"**建议：** {issue.recommendation}",
                "",
            ]
        )
    lines.extend(["## 收入变化驱动", "", present_drivers(drivers).to_markdown(index=False), ""])
    if targets is not None and not targets.empty:
        lines.extend(["## 目标状态", "", targets.to_markdown(index=False), ""])
    if forecast is not None and not forecast.empty:
        lines.extend(["## 未来预测", "", forecast.to_markdown(index=False), ""])
    return "\n".join(lines)


def render_html_report(
    period: str,
    kpi: dict[str, float],
    issues: list[DiagnosticIssue],
    drivers: pd.DataFrame,
    trend: pd.DataFrame,
    *,
    profit_table: pd.DataFrame | None = None,
    targets: pd.DataFrame | None = None,
    forecast: pd.DataFrame | None = None,
    inventory: pd.DataFrame | None = None,
    quality: pd.DataFrame | None = None,
    data_profile: str | None = None,
) -> str:
    summary, engine = executive_summary(issues)
    return HTML_TEMPLATE.render(
        period=period,
        generated_at=datetime.now().strftime("%Y-%m-%d %H:%M"),
        summary_engine=engine,
        summary=summary,
        data_profile=data_profile or "未提供来源标签",
        kpis=format_kpis(kpi),
        issues=issues,
        drivers_table=_safe_table(present_drivers(drivers)),
        profit_table=_safe_table(profit_table),
        targets_table=_safe_table(present_targets(targets)),
        forecast_table=_safe_table(present_forecast(forecast)),
        inventory_table=_safe_table(present_inventory(inventory)),
        quality_table=_safe_table(present_quality(quality)),
        trend_table=_safe_table(trend.tail(12)),
    )


def build_html_report(
    output_path: str | Path,
    period: str,
    kpi: dict[str, float],
    issues: list[DiagnosticIssue],
    drivers: pd.DataFrame,
    trend: pd.DataFrame,
    **kwargs,
) -> Path:
    html = render_html_report(period, kpi, issues, drivers, trend, **kwargs)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(html, encoding="utf-8")
    return output


def _add_df_table(doc: Document, df: pd.DataFrame, *, max_rows: int = 20) -> None:
    shown = df.head(max_rows).copy()
    if shown.empty:
        doc.add_paragraph("暂无数据。")
        return
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(shown.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                cells[idx].text = f"{value:,.3f}"
            else:
                cells[idx].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _create_word_report(
    period: str,
    kpi: dict[str, float],
    issues: list[DiagnosticIssue],
    drivers: pd.DataFrame,
    *,
    targets: pd.DataFrame | None = None,
    forecast: pd.DataFrame | None = None,
    inventory: pd.DataFrame | None = None,
    quality: pd.DataFrame | None = None,
    data_profile: str | None = None,
) -> Document:
    summary, engine = executive_summary(issues)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    title = doc.add_heading("InsightFlow 经营分析与决策报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(f"分析期间：{period}｜生成时间：{datetime.now():%Y-%m-%d %H:%M}｜摘要引擎：{engine}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trust = doc.add_paragraph(f"数据状态：{data_profile or '未提供来源标签'}")
    trust.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、管理层摘要", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("二、核心经营指标", level=1)
    kpi_table = doc.add_table(rows=1, cols=3)
    kpi_table.style = "Table Grid"
    kpi_table.rows[0].cells[0].text = "指标"
    kpi_table.rows[0].cells[1].text = "数值"
    kpi_table.rows[0].cells[2].text = "经营含义"
    meanings = {
        "净销售额": "扣除折扣后的有效销售收入",
        "毛利润": "净销售额扣除商品成本",
        "毛利率": "商品结构与采购成本质量",
        "贡献利润": "进一步扣除物流、支付和营销费用",
        "贡献利润率": "规模增长是否真正创造利润",
        "订单量": "交易规模",
        "活跃客户": "产生有效购买的客户数",
        "复购率": "客户质量与收入可持续性",
        "取消率": "履约和商品质量风险",
    }
    for label, value in format_kpis(kpi):
        cells = kpi_table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[2].text = meanings.get(label, "")

    doc.add_heading("三、关键诊断与行动建议", level=1)
    for issue in issues:
        doc.add_heading(f"{issue.severity}｜{issue.title}", level=2)
        doc.add_paragraph(issue.finding)
        doc.add_paragraph(f"证据：{issue.evidence}")
        doc.add_paragraph(f"建议：{issue.recommendation}")

    doc.add_heading("四、收入变化驱动", level=1)
    _add_df_table(doc, present_drivers(drivers))

    if targets is not None:
        doc.add_heading("五、目标完成状态", level=1)
        _add_df_table(doc, present_targets(targets))
    if forecast is not None:
        doc.add_heading("六、未来预测", level=1)
        _add_df_table(doc, present_forecast(forecast))
    if inventory is not None:
        doc.add_heading("七、库存风险", level=1)
        _add_df_table(doc, present_inventory(inventory), max_rows=15)
    if quality is not None:
        doc.add_heading("八、数据质量", level=1)
        _add_df_table(doc, present_quality(quality))

    doc.add_heading("九、管理行动清单", level=1)
    for issue in issues:
        doc.add_paragraph(issue.recommendation, style="List Bullet")

    doc.add_heading("十、核心指标口径", level=1)
    definitions = [
        "净销售额：有效正向交易商品毛额减折扣金额。",
        "毛利润：净销售额减商品成本。",
        "贡献利润：毛利润减物流、支付手续费、营销费用和退货处理成本。",
        "复购率：期间内订单数不少于2笔的客户占购买客户比例。",
        "收入驱动：采用Shapley方法，将收入变化分解为活跃客户、购买频次和客单价贡献。",
        "预测：在多个基线模型中通过滚动回测选择误差最低模型，并给出不确定区间。",
    ]
    for item in definitions:
        doc.add_paragraph(item, style="List Bullet")

    return doc


def build_word_report(
    output_path: str | Path,
    period: str,
    kpi: dict[str, float],
    issues: list[DiagnosticIssue],
    drivers: pd.DataFrame,
    **kwargs,
) -> Path:
    doc = _create_word_report(period, kpi, issues, drivers, **kwargs)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def build_word_report_bytes(
    period: str,
    kpi: dict[str, float],
    issues: list[DiagnosticIssue],
    drivers: pd.DataFrame,
    **kwargs,
) -> bytes:
    doc = _create_word_report(period, kpi, issues, drivers, **kwargs)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
